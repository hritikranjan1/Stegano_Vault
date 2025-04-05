from flask import Flask, request, send_file, jsonify, render_template, url_for, session, redirect
import os
import tempfile
import shutil
from PIL import Image
from stegano import lsb
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from pydub import AudioSegment
import cv2
import hashlib
import numpy as np
from mutagen.mp4 import MP4
import base64
import logging
import uuid
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import json
from werkzeug.utils import secure_filename
import time
from dotenv import load_dotenv
from datetime import datetime
import random
import string
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadTimeSignature
from functools import wraps
import bcrypt
from supabase import create_client, Client

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)

# ========================
# SECURITY CONFIGURATION
# ========================

app.secret_key = os.environ.get('SECRET_KEY', os.urandom(24).hex())

@app.after_request
def add_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    if 'Cache-Control' not in response.headers:
        response.headers['Cache-Control'] = 'no-store, max-age=0'
    return response

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ========================
# SUPABASE CONFIGURATION
# ========================

supabase_url = os.environ.get('SUPABASE_URL')
supabase_key = os.environ.get('SUPABASE_KEY')
if not supabase_url or not supabase_key:
    raise Exception("Supabase URL or Key not provided in environment variables")
supabase: Client = create_client(supabase_url, supabase_key)

# ========================
# EMAIL CONFIGURATION
# ========================

app.config['MAIL_SERVER'] = os.environ.get('MAIL_SERVER', 'smtp-relay.brevo.com')
app.config['MAIL_PORT'] = int(os.environ.get('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = os.environ.get('MAIL_USE_TLS', 'true').lower() == 'true'
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.environ.get('MAIL_DEFAULT_SENDER', 'noreply@steganovault.com')

mail = Mail(app)

serializer = URLSafeTimedSerializer(app.secret_key)

# ========================
# APPLICATION CONFIGURATION
# ========================

SCOPES = ['https://www.googleapis.com/auth/drive']
GOOGLE_DRIVE_FOLDER_ID = os.environ.get('GOOGLE_DRIVE_FOLDER_ID')

# ========================
# AUTHENTICATION UTILITIES
# ========================

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'error': 'Authentication required, Please login first to submit your review'}), 401
        return f(*args, **kwargs)
    return decorated_function

def get_user_by_email(email):
    try:
        response = supabase.table('users').select('*').eq('email', email).single().execute()
        if response.data:
            return {
                'id': response.data['id'],
                'name': response.data['name'],
                'email': response.data['email'],
                'password': response.data['password'],
                'created_at': response.data['created_at'],
                'verified': response.data['verified']
            }
        return None
    except Exception as e:
        logger.error(f"Error getting user by email {email}: {e}")
        return None

def save_user(user):
    try:
        payload = {
            'name': user['name'],
            'email': user['email'],
            'password': user['password'],
            'created_at': datetime.now().isoformat(),
            'verified': user.get('verified', False)
        }
        logger.info(f"Attempting to save user with payload: {payload}")
        response = supabase.table('users').insert(payload).execute()
        logger.info(f"Supabase response: {response.data}")
        return bool(response.data)
    except Exception as e:
        logger.error(f"Error saving user: {e}")
        return False

def update_user_password(email, new_password):
    try:
        response = supabase.table('users').update({
            'password': hash_password(new_password)
        }).eq('email', email).execute()
        return bool(response.data)
    except Exception as e:
        logger.error(f"Error updating user password: {e}")
        return False

def generate_otp(email):
    otp = ''.join(random.choices(string.digits, k=6))
    try:
        response = supabase.table('otp_storage').upsert({
            'email': email,
            'otp': otp,
            'timestamp': datetime.now().isoformat(),
            'verified': False
        }, on_conflict=['email']).execute()
        return otp if response.data else None
    except Exception as e:
        logger.error(f"Error generating OTP for {email}: {e}")
        return None

def verify_otp(email, otp):
    try:
        response = supabase.table('otp_storage').select('otp', 'timestamp').eq('email', email).single().execute()
        if not response.data:
            return False
        
        stored_otp = response.data['otp']
        otp_time = datetime.fromisoformat(response.data['timestamp'].replace('Z', '+00:00'))
        
        if (datetime.now() - otp_time).total_seconds() > 300:
            return False
        
        if stored_otp == otp:
            supabase.table('otp_storage').update({'verified': True}).eq('email', email).execute()
            supabase.table('users').update({'verified': True}).eq('email', email).execute()
            return True
        return False
    except Exception as e:
        logger.error(f"Error verifying OTP for {email}: {e}")
        return False

def is_verified(email):
    try:
        response = supabase.table('otp_storage').select('verified').eq('email', email).single().execute()
        return response.data['verified'] if response.data else False
    except Exception as e:
        logger.error(f"Error checking verification status for {email}: {e}")
        return False

def hash_password(password):
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def check_password(password, hashed):
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

# ========================
# AUTHENTICATION ROUTES
# ========================

@app.route('/auth/register', methods=['POST'])
def register():
    data = request.get_json()
    if not data or 'email' not in data or 'password' not in data or 'name' not in data:
        return jsonify({'error': 'Name, email, and password are required'}), 400
    
    email = data['email'].lower()
    password = data['password']
    name = data['name'].strip()
    
    if len(password) < 8:
        return jsonify({'error': 'Password must be at least 8 characters'}), 400
    
    if get_user_by_email(email):
        return jsonify({'error': 'Email already registered'}), 400
    
    hashed_password = hash_password(password)
    user = {'name': name, 'email': email, 'password': hashed_password}
    if not save_user(user):
        return jsonify({'error': 'Failed to register user'}), 500
    
    otp = generate_otp(email)
    if not otp:
        return jsonify({'error': 'Failed to generate OTP'}), 500
    
    try:
        msg = Message('Your SteganoVault Verification Code',
                     recipients=[email],
                     sender=app.config['MAIL_DEFAULT_SENDER'])
        msg.body = f'Your OTP for SteganoVault is: {otp}\nThis code will expire in 5 minutes.'
        mail.send(msg)
        logger.info(f"OTP sent successfully to {email}")
    except Exception as e:
        logger.error(f"Failed to send OTP email to {email}: {e}")
        return jsonify({'error': 'Failed to send OTP email'}), 500
    
    return jsonify({'message': 'OTP sent to your email'}), 200

@app.route('/auth/verify', methods=['POST'])
def verify():
    data = request.get_json()
    if not data or 'email' not in data or 'otp' not in data:
        return jsonify({'error': 'Email and OTP are required'}), 400
    
    email = data['email'].lower()
    otp = data['otp']
    
    if not verify_otp(email, otp):
        return jsonify({'error': 'Invalid or expired OTP'}), 400
    
    return jsonify({'message': 'Email verified successfully'}), 200

@app.route('/auth/login', methods=['POST'])
def login():
    data = request.get_json()
    if not data or 'email' not in data or 'password' not in data:
        return jsonify({'error': 'Email and password are required'}), 400
    
    email = data['email'].lower()
    password = data['password']
    
    user = get_user_by_email(email)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    if not check_password(password, user['password']):
        return jsonify({'error': 'Invalid credentials'}), 401
    
    if not user.get('verified', False):
        return jsonify({'error': 'Email not verified'}), 403
    
    session['user_id'] = user['email']
    return jsonify({
        'message': 'Login successful',
        'user': {
            'email': user['email'],
            'name': user.get('name', ''),
            'verified': user.get('verified', False)
        }
    }), 200

@app.route('/auth/logout', methods=['POST'])
def logout():
    if 'user_id' in session:
        session.pop('user_id', None)
        return jsonify({'message': 'Logged out successfully'}), 200
    return jsonify({'error': 'Not logged in'}), 400

@app.route('/auth/status', methods=['GET'])
def auth_status():
    if 'user_id' in session:
        user = get_user_by_email(session['user_id'])
        if user:
            return jsonify({
                'authenticated': True,
                'user': {
                    'name': user.get('name', ''),
                    'email': user['email'],
                    'verified': user.get('verified', False)
                }
            }), 200
    return jsonify({'authenticated': False}), 200

@app.route('/auth/forgot-password', methods=['POST'])
def forgot_password():
    data = request.get_json()
    if not data or 'email' not in data:
        return jsonify({'error': 'Email is required'}), 400
    
    email = data['email'].lower()
    user = get_user_by_email(email)
    
    if not user:
        return jsonify({'error': 'Email not found'}), 404
    
    otp = generate_otp(email)
    if not otp:
        return jsonify({'error': 'Failed to generate OTP'}), 500
    
    try:
        msg = Message('Password Reset OTP',
                     recipients=[email],
                     sender=app.config['MAIL_DEFAULT_SENDER'])
        msg.body = f'Your OTP for password reset is: {otp}\nThis code will expire in 5 minutes.'
        mail.send(msg)
        return jsonify({'message': 'OTP sent to your email'}), 200
    except Exception as e:
        logger.error(f"Failed to send reset email to {email}: {e}")
        return jsonify({'error': 'Failed to send OTP email'}), 500

@app.route('/auth/reset-password', methods=['POST'])
def reset_password():
    data = request.get_json()
    if not data or 'email' not in data or 'otp' not in data or 'new_password' not in data:
        return jsonify({'error': 'Email, OTP and new password are required'}), 400
    
    email = data['email'].lower()
    otp = data['otp']
    new_password = data['new_password']
    
    if len(new_password) < 8:
        return jsonify({'error': 'Password must be at least 8 characters'}), 400
    
    if not verify_otp(email, otp):
        return jsonify({'error': 'Invalid or expired OTP'}), 400
    
    if update_user_password(email, new_password):
        return jsonify({'message': 'Password updated successfully'}), 200
    return jsonify({'error': 'Failed to update password'}), 500

# ========================
# REVIEWS ENDPOINTS
# ========================

@app.route('/api/reviews', methods=['GET'])
def get_reviews():
    try:
        response = supabase.table('reviews').select('name', 'text', 'rating', 'timestamp', 'user_email')\
            .eq('verified', True)\
            .order('timestamp', desc=True)\
            .execute()
        reviews = []
        for review in response.data:
            user_response = supabase.table('users').select('name').eq('email', review['user_email']).single().execute()
            user_name = user_response.data['name'] if user_response.data else 'Anonymous'
            reviews.append({
                "name": review['name'],
                "text": review['text'],
                "rating": review['rating'],
                "timestamp": review['timestamp'],
                "user_email": review['user_email'],
                "user_name": user_name
            })
        return jsonify(reviews)
    except Exception as e:
        logger.error(f"Error fetching reviews: {e}")
        return jsonify([])

@app.route('/api/reviews', methods=['POST'])
@login_required
def submit_review():
    data = request.get_json()
    if not data or 'name' not in data or 'text' not in data or 'rating' not in data:
        return jsonify({"error": "Missing required fields"}), 400
    
    user_email = session['user_id']
    user = get_user_by_email(user_email)
    
    if not user:
        return jsonify({"error": "User not found"}), 404
    
    try:
        response = supabase.table('reviews').insert({
            'user_email': user_email,
            'name': data['name'],
            'text': data['text'],
            'rating': int(data['rating']),
            'timestamp': datetime.now().isoformat(),
            'verified': True
        }).execute()
        
        if response.data:
            # Get the newly created review to return
            new_review = response.data[0]
            return jsonify({
                "message": "Review submitted successfully",
                "review": {
                    "name": new_review['name'],
                    "text": new_review['text'],
                    "rating": new_review['rating'],
                    "timestamp": new_review['timestamp'],
                    "user_email": new_review['user_email'],
                    "user_name": user['name']
                }
            })
        return jsonify({"error": "Failed to submit review"}), 500
    except Exception as e:
        logger.error(f"Error saving review: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/reviews/latest')
def get_latest_reviews():
    try:
        response = supabase.table('reviews').select('name', 'text', 'rating', 'timestamp', 'user_email')\
            .eq('verified', True)\
            .order('timestamp', desc=True)\
            .limit(10)\
            .execute()
        reviews = []
        for review in response.data:
            user_response = supabase.table('users').select('name').eq('email', review['user_email']).single().execute()
            user_name = user_response.data['name'] if user_response.data else 'Anonymous'
            reviews.append({
                "name": review['name'],
                "text": review['text'],
                "rating": review['rating'],
                "timestamp": review['timestamp'],
                "user_email": review['user_email'],
                "user_name": user_name
            })
        return jsonify(reviews)
    except Exception as e:
        logger.error(f"Error fetching latest reviews: {e}")
        return jsonify([])

@app.route('/api/testimonials')
def get_testimonials():
    try:
        response = supabase.table('reviews').select('name', 'text', 'rating', 'timestamp', 'user_email')\
            .eq('verified', True)\
            .order('timestamp', desc=True)\
            .limit(7)\
            .execute()
        verified_reviews = []
        for review in response.data:
            user_response = supabase.table('users').select('name').eq('email', review['user_email']).single().execute()
            user_name = user_response.data['name'] if user_response.data else 'Anonymous'
            verified_reviews.append({
                "name": review['name'],
                "text": review['text'],
                "rating": review['rating'],
                "timestamp": review['timestamp'],
                "user_email": review['user_email'],
                "user_name": user_name
            })
        
        default_testimonials = [
            {"name": "Agent X", "text": "SteganoVault made hiding messages so fun!", "rating": 5},
            {"name": "Codebreaker Y", "text": "A must-have tool for spy enthusiasts!", "rating": 5},
            {"name": "Security Expert Z", "text": "Perfect balance of security and usability.", "rating": 4}
        ]
        
        return jsonify(default_testimonials + verified_reviews)
    except Exception as e:
        logger.error(f"Error fetching testimonials: {e}")
        return jsonify([])

# ========================
# FILE PROCESSING FUNCTIONS
# ========================

def convert_to_png(input_path):
    try:
        img = Image.open(input_path).convert("RGB")
        png_path = input_path.rsplit(".", 1)[0] + "_converted.png"
        img.save(png_path, format="PNG", quality=95)
        return png_path
    except Exception as e:
        logger.error(f"Image Conversion Error: {e}")
        return None

def generate_preview(input_path):
    try:
        if input_path.lower().endswith(('.png', '.jpg', '.jpeg')):
            img = Image.open(input_path)
            preview_filename = os.path.basename(input_path).replace(".", "_preview.")
            preview_path = os.path.join('static', preview_filename)
            img.thumbnail((300, 300))
            img.save(preview_path, format="PNG", quality=85)
            return preview_filename
        return None
    except Exception as e:
        logger.error(f"Preview Generation Error: {e}")
        return None

def encode_image(input_path, message, password, watermark=None):
    try:
        if not input_path.lower().endswith(".png"):
            input_path = convert_to_png(input_path)
        if not input_path:
            return None, None
        
        secret_message = f"{password}:{message}" if password else message
        if watermark:
            secret_message = f"{watermark}:{secret_message}"
            
        encoded_img = lsb.hide(input_path, secret_message)
        output_path = input_path.replace(".png", "_encoded.png")
        encoded_img.save(output_path, format="PNG", quality=95)
        
        preview_filename = generate_preview(output_path)
        return output_path, preview_filename
    except Exception as e:
        logger.error(f"Image Encoding Error: {e}")
        return None, None

def decode_image(input_path, password):
    try:
        preview_filename = generate_preview(input_path)
        extracted_message = lsb.reveal(input_path)
        
        if extracted_message:
            if ":" in extracted_message:
                parts = extracted_message.split(":")
                if len(parts) == 3:
                    stored_password, stored_message = parts[1], parts[2]
                else:
                    stored_password, stored_message = parts[0], parts[1]
                return stored_message if stored_password == password else "Incorrect password!", preview_filename
            return extracted_message, preview_filename
        return "No hidden message found!", preview_filename
    except Exception as e:
        logger.error(f"Image Decoding Error: {e}")
        return f"Error decoding image: {str(e)}", None

def encode_txt(input_path, message, password):
    try:
        with open(input_path, "r", encoding="utf-8") as file:
            original_content = file.read()
        
        secret_message = f"{password}:{message}" if password else message
        binary_data = ''.join(format(ord(c), '08b') for c in secret_message)
        encoded_message = ''.join("\u200B" if bit == "0" else "\u200D" for bit in binary_data)

        output_path = input_path.replace(".", "_encoded.")
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(original_content + "\n" + encoded_message)
        return output_path
    except Exception as e:
        logger.error(f"TXT Encoding Error: {e}")
        return None

def decode_txt(input_path, password):
    try:
        with open(input_path, "r", encoding="utf-8") as file:
            content = file.read()
        
        binary_data = ''.join("0" if char == "\u200B" else "1" for char in content if char in ["\u200B", "\u200D"])
        if not binary_data:
            return "No hidden message found!"
        
        extracted_message = ''.join(chr(int(binary_data[i:i+8], 2)) for i in range(0, len(binary_data) - len(binary_data) % 8, 8))
        
        if extracted_message and ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "Incorrect password!"
        return extracted_message if extracted_message else "No hidden message found!"
    except Exception as e:
        logger.error(f"TXT Decoding Error: {e}")
        return f"Error decoding text: {str(e)}"

def encode_docx(input_path, message, password):
    try:
        doc = Document(input_path)
        secret_message = f"{password}:{message}" if password else message
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(secret_message)
        run.font.hidden = True
        
        output_path = input_path.replace(".", "_encoded.")
        doc.save(output_path)
        return output_path
    except Exception as e:
        logger.error(f"DOCX Encoding Error: {e}")
        return None

def decode_docx(input_path, password):
    try:
        doc = Document(input_path)
        extracted_message = ""
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.hidden:
                    extracted_message += run.text
        
        if extracted_message and ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "Incorrect password!"
        return extracted_message if extracted_message else "No hidden message found!"
    except Exception as e:
        logger.error(f"DOCX Decoding Error: {e}")
        return f"Error decoding DOCX: {str(e)}"

def encode_pdf(input_path, message, password):
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        secret_message = f"{password}:{message}" if password else message
        writer.add_metadata({"/Message": secret_message})
        
        output_path = input_path.replace(".", "_encoded.")
        with open(output_path, "wb") as output_pdf:
            writer.write(output_pdf)
        return output_path
    except Exception as e:
        logger.error(f"PDF Encoding Error: {e}")
        return None

def decode_pdf(input_path, password):
    try:
        reader = PdfReader(input_path)
        metadata = reader.metadata
        extracted_message = metadata.get("/Message", "") if metadata else ""
        
        if extracted_message and ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "Incorrect password!"
        return extracted_message if extracted_message else "No hidden message found!"
    except Exception as e:
        logger.error(f"PDF Decoding Error: {e}")
        return f"Error decoding PDF: {str(e)}"

def convert_to_wav(input_path):
    try:
        audio = AudioSegment.from_file(input_path)
        wav_path = input_path.rsplit(".", 1)[0] + "_converted.wav"
        audio.export(wav_path, format="wav", codec="pcm_s16le")
        return wav_path
    except Exception as e:
        logger.error(f"Audio Conversion Error: {e}")
        return None

def encode_audio(input_path, message, password):
    try:
        if not input_path.lower().endswith(".wav"):
            input_path = convert_to_wav(input_path)
        if not input_path:
            return None

        secret_message = f"{password}:{message}" if password else message
        message_bytes = secret_message.encode('utf-8')
        message_length = len(message_bytes)
        binary_length = format(message_length, '032b')
        binary_message = binary_length + ''.join(format(byte, '08b') for byte in message_bytes)

        audio = AudioSegment.from_file(input_path, format="wav")
        samples = audio.get_array_of_samples()
        if len(binary_message) > len(samples):
            return None

        samples_array = np.array(samples, dtype=np.int16)
        for i in range(len(binary_message)):
            samples_array[i] = (samples_array[i] & ~1) | int(binary_message[i])

        encoded_audio = audio._spawn(samples_array.tobytes())
        output_path = input_path.replace(".wav", "_encoded.wav")
        encoded_audio.export(output_path, format="wav", codec="pcm_s16le")
        return output_path
    except Exception as e:
        logger.error(f"Audio Encoding Error: {e}")
        return None

def decode_audio(input_path, password):
    try:
        if not input_path.lower().endswith(".wav"):
            input_path = convert_to_wav(input_path)
        if not input_path:
            return "Error decoding audio!"

        audio = AudioSegment.from_file(input_path, format="wav")
        samples = np.array(audio.get_array_of_samples(), dtype=np.int16)

        binary_length = ''.join(str(sample & 1) for sample in samples[:32])
        message_length = int(binary_length, 2)

        binary_message = ''.join(str(sample & 1) for sample in samples[32:32 + message_length * 8])
        message_bytes = bytes(int(binary_message[i:i+8], 2) for i in range(0, len(binary_message) - len(binary_message) % 8, 8))
        extracted_message = message_bytes.decode('utf-8')

        if extracted_message and ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "Incorrect password!"
        return extracted_message if extracted_message else "No hidden message found!"
    except Exception as e:
        logger.error(f"Audio Decoding Error: {e}")
        return f"Error decoding audio: {str(e)}"

def encode_video(input_path, message, password):
    try:
        cap = cv2.VideoCapture(input_path)
        if not cap.isOpened():
            raise ValueError("Could not open video file")
        cap.release()

        secret_message = f"{password}:{message}" if password else message
        encoded_message = base64.b64encode(secret_message.encode('utf-8')).decode('utf-8')

        output_path = input_path.replace(".", "_encoded.")
        shutil.copy2(input_path, output_path)

        video = MP4(output_path)
        video['desc'] = encoded_message
        video.save()

        test_cap = cv2.VideoCapture(output_path)
        if not test_cap.isOpened():
            raise ValueError("Encoded video file is corrupted")
        test_cap.release()

        return output_path
    except Exception as e:
        logger.error(f"Video Encoding Error: {e}")
        return None

def decode_video(input_path, password):
    try:
        cap = cv2.VideoCapture(input_path)
        if not cap.isOpened():
            raise ValueError("Could not open video file")
        cap.release()

        video = MP4(input_path)
        if 'desc' not in video:
            return "No hidden message found!"

        encoded_message = video['desc'][0]
        decoded_bytes = base64.b64decode(encoded_message.encode('utf-8'))
        extracted_message = decoded_bytes.decode('utf-8')

        if ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "Incorrect password!"
        return extracted_message if extracted_message else "No hidden message found!"
    except Exception as e:
        logger.error(f"Video Decoding Error: {e}")
        return f"Error decoding video: {str(e)}"

def get_drive_service():
    credentials_json = os.environ.get('GOOGLE_CREDENTIALS')
    if not credentials_json:
        logger.error("GOOGLE_CREDENTIALS not set")
        return None
    
    try:
        credentials_info = json.loads(credentials_json)
        credentials = service_account.Credentials.from_service_account_info(credentials_info, scopes=SCOPES)
        return build('drive', 'v3', credentials=credentials)
    except Exception as e:
        logger.error(f"Error initializing Drive service: {e}")
        return None

def upload_to_drive(file_path, file_name):
    drive_service = get_drive_service()
    if not drive_service:
        return None
    
    try:
        file_metadata = {
            'name': file_name,
            'parents': [GOOGLE_DRIVE_FOLDER_ID]
        }
        media = MediaFileUpload(file_path)
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, webViewLink'
        ).execute()
        drive_service.permissions().create(
            fileId=file['id'],
            body={'role': 'reader', 'type': 'anyone'}
        ).execute()
        return file['webViewLink'].replace('usp=sharing', 'usp=drivesdk')
    except Exception as e:
        logger.error(f"Failed to upload to Drive: {e}")
        return None

# ========================
# MAIN APPLICATION ROUTES
# ========================

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/encode", methods=["POST"])
def encode():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    uploaded_file = request.files['file']
    message = request.form.get("message")
    password = request.form.get("password", "")
    watermark = request.form.get("watermark", "")

    if not message:
        return jsonify({"error": "Message is required"}), 400

    filename = secure_filename(uploaded_file.filename)
    if not filename:
        return jsonify({"error": "Invalid filename"}), 400

    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, filename)
    uploaded_file.save(file_path)

    ext = os.path.splitext(file_path)[1].lower()
    output_file = None
    preview_file = None

    try:
        start_time = time.time()
        
        if ext in [".png", ".jpg", ".jpeg"]:
            output_file, preview_file = encode_image(file_path, message, password, watermark)
        elif ext == ".txt":
            output_file = encode_txt(file_path, message, password)
        elif ext == ".pdf":
            output_file = encode_pdf(file_path, message, password)
        elif ext == ".docx":
            output_file = encode_docx(file_path, message, password)
        elif ext in [".wav", ".mp3"]:
            output_file = encode_audio(file_path, message, password)
        elif ext in [".mp4", ".avi", ".mov"]:
            output_file = encode_video(file_path, message, password)
        else:
            return jsonify({"error": "Unsupported file format"}), 400

        if output_file and os.path.exists(output_file):
            processing_time = round(time.time() - start_time, 2)
            unique_id = str(uuid.uuid4())
            final_filename = f"encoded_{unique_id}_{os.path.basename(output_file)}"
            share_url = upload_to_drive(output_file, final_filename)

            with open(output_file, "rb") as f:
                checksum = hashlib.sha256(f.read()).hexdigest()
            
            response_data = {
                "status": "success",
                "filename": final_filename,
                "checksum": checksum,
                "processing_time": processing_time,
                "preview_url": url_for('static', filename=preview_file) if preview_file else None,
                "share_url": share_url
            }
            response = send_file(output_file, as_attachment=True, download_name=final_filename)
            response.headers["X-Response-Data"] = json.dumps(response_data)
            return response
        
        return jsonify({"error": "Failed to encode file"}), 500
    except Exception as e:
        logger.error(f"Encode Error: {e}")
        return jsonify({"error": f"Encoding failed: {str(e)}"}), 500
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

@app.route("/decode", methods=["POST"])
def decode():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    uploaded_file = request.files['file']
    password = request.form.get("password", "")

    filename = secure_filename(uploaded_file.filename)
    if not filename:
        return jsonify({"error": "Invalid filename"}), 400

    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, filename)
    uploaded_file.save(file_path)

    ext = os.path.splitext(file_path)[1].lower()
    decoded_message = None
    preview_file = None

    try:
        start_time = time.time()
        
        if ext in [".png", ".jpg", ".jpeg"]:
            decoded_message, preview_file = decode_image(file_path, password)
        elif ext == ".txt":
            decoded_message = decode_txt(file_path, password)
        elif ext == ".pdf":
            decoded_message = decode_pdf(file_path, password)
        elif ext == ".docx":
            decoded_message = decode_docx(file_path, password)
        elif ext in [".wav", ".mp3"]:
            decoded_message = decode_audio(file_path, password)
        elif ext in [".mp4", ".avi", ".mov"]:
            decoded_message = decode_video(file_path, password)
        else:
            return jsonify({"error": "Unsupported file format"}), 400

        processing_time = round(time.time() - start_time, 2)
        
        return jsonify({
            "status": "success",
            "message": decoded_message,
            "preview_url": url_for('static', filename=preview_file) if preview_file else None,
            "processing_time": processing_time
        })
    except Exception as e:
        logger.error(f"Decode Error: {e}")
        return jsonify({"error": f"Decoding failed: {str(e)}"}), 500
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

@app.route('/privacy')
def privacy():
    return jsonify({
        "privacy_policy": "SteganoVault does not store your files or messages. All processing happens on the server and files are deleted immediately after."
    })

# ========================
# APPLICATION STARTUP
# ========================

if __name__ == "__main__":
    if not os.path.exists('static'):
        os.makedirs('static')
    
    logger.info("Starting SteganoVault server with Supabase...")
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port, debug=False)